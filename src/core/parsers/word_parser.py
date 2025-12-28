"""Word document parser using python-docx."""

from pathlib import Path
from typing import Optional
import io

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

from src.core.parsers.base import (
    BaseParser,
    ContentType,
    ParsedDocument,
    ParsedSection,
)


class WordParser(BaseParser):
    """Parser for Word (.docx) files."""

    SUPPORTED_EXTENSIONS = {".docx"}

    # Heading style mappings
    HEADING_STYLES = {
        "Heading 1": 1,
        "Heading 2": 2,
        "Heading 3": 3,
        "Heading 4": 4,
        "Heading 5": 5,
        "Heading 6": 6,
        "Title": 1,
    }

    def parse(self, file_path: Path) -> ParsedDocument:
        """Parse a Word document."""
        doc = Document(file_path)
        return self._parse_document(doc, file_path.name, file_path)

    def parse_bytes(self, content: bytes, filename: str) -> ParsedDocument:
        """Parse Word document from bytes."""
        doc = Document(io.BytesIO(content))
        return self._parse_document(doc, filename, Path(filename))

    def _parse_document(
        self,
        doc: Document,
        filename: str,
        file_path: Path,
    ) -> ParsedDocument:
        """Parse Word document into sections."""
        sections: list[ParsedSection] = []
        current_hierarchy: list[str] = []
        current_text_parts: list[str] = []
        current_heading: Optional[str] = None
        current_level: Optional[int] = None

        for element in doc.element.body:
            # Handle paragraphs
            if element.tag.endswith("p"):
                para = Paragraph(element, doc)
                style_name = para.style.name if para.style else ""
                text = para.text.strip()

                if not text:
                    continue

                # Check if it's a heading
                heading_level = self.HEADING_STYLES.get(style_name)
                if heading_level:
                    # Save accumulated text
                    if current_text_parts:
                        sections.append(
                            ParsedSection(
                                content="\n".join(current_text_parts),
                                content_type=ContentType.TEXT,
                                heading_text=current_heading,
                                heading_level=current_level,
                                section_hierarchy=current_hierarchy.copy(),
                            )
                        )
                        current_text_parts = []

                    # Update hierarchy
                    while len(current_hierarchy) >= heading_level:
                        current_hierarchy.pop()
                    current_hierarchy.append(text)

                    current_heading = text
                    current_level = heading_level

                    # Add heading section
                    sections.append(
                        ParsedSection(
                            content=text,
                            content_type=ContentType.HEADING,
                            heading_level=heading_level,
                            heading_text=text,
                            section_hierarchy=current_hierarchy.copy(),
                        )
                    )
                else:
                    # Check for list items
                    if self._is_list_paragraph(para):
                        current_text_parts.append(f"- {text}")
                    else:
                        current_text_parts.append(text)

            # Handle tables
            elif element.tag.endswith("tbl"):
                # Save accumulated text first
                if current_text_parts:
                    sections.append(
                        ParsedSection(
                            content="\n".join(current_text_parts),
                            content_type=ContentType.TEXT,
                            heading_text=current_heading,
                            heading_level=current_level,
                            section_hierarchy=current_hierarchy.copy(),
                        )
                    )
                    current_text_parts = []

                # Parse table
                table = Table(element, doc)
                table_content = self._parse_table(table)

                sections.append(
                    ParsedSection(
                        content=table_content,
                        content_type=ContentType.TABLE,
                        heading_text=current_heading,
                        heading_level=current_level,
                        section_hierarchy=current_hierarchy.copy(),
                    )
                )

        # Don't forget remaining text
        if current_text_parts:
            sections.append(
                ParsedSection(
                    content="\n".join(current_text_parts),
                    content_type=ContentType.TEXT,
                    heading_text=current_heading,
                    heading_level=current_level,
                    section_hierarchy=current_hierarchy.copy(),
                )
            )

        return ParsedDocument(
            filename=filename,
            file_path=file_path,
            sections=sections,
            metadata=self._extract_metadata(doc),
        )

    def _is_list_paragraph(self, para: Paragraph) -> bool:
        """Check if paragraph is a list item."""
        try:
            # Check for numbering
            pPr = para._p.pPr
            if pPr is not None:
                numPr = pPr.numPr
                return numPr is not None
        except Exception:
            pass
        return False

    def _parse_table(self, table: Table) -> str:
        """Convert Word table to Markdown format."""
        rows = []
        for i, row in enumerate(table.rows):
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            row_str = "| " + " | ".join(cells) + " |"
            rows.append(row_str)

            # Add header separator after first row
            if i == 0:
                separator = "| " + " | ".join(["---"] * len(cells)) + " |"
                rows.append(separator)

        return "\n".join(rows)

    def _extract_metadata(self, doc: Document) -> dict:
        """Extract metadata from Word document."""
        metadata = {}

        # Core properties
        if doc.core_properties:
            props = doc.core_properties
            if props.title:
                metadata["title"] = props.title
            if props.author:
                metadata["author"] = props.author
            if props.subject:
                metadata["subject"] = props.subject
            if props.keywords:
                metadata["keywords"] = props.keywords
            if props.created:
                metadata["created"] = props.created.isoformat()
            if props.modified:
                metadata["modified"] = props.modified.isoformat()

        # Count elements
        metadata["paragraph_count"] = len(doc.paragraphs)
        metadata["table_count"] = len(doc.tables)

        return metadata
